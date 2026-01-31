#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

import logging
import copy
import re
import os

from common.constants import ParserType
from io import BytesIO
from rag.nlp import rag_tokenizer, tokenize, tokenize_table, bullets_category, title_frequency, tokenize_chunks, docx_question_level, attach_media_context
from common.token_utils import num_tokens_from_string
from deepdoc.parser import PdfParser, DocxParser
from deepdoc.parser.figure_parser import vision_figure_parser_pdf_wrapper, vision_figure_parser_docx_wrapper
from docx import Document
from PIL import Image
from rag.app.naive import by_plaintext, PARSERS
from common.parser_config_utils import normalize_layout_recognizer


class Pdf(PdfParser):
    def __init__(self):
        self.model_speciess = ParserType.MANUAL.value
        super().__init__()

    def __call__(self, filename, binary=None, from_page=0, to_page=100000, zoomin=3, callback=None):
        from timeit import default_timer as timer

        start = timer()
        callback(msg="OCR started")
        self.__images__(filename if not binary else binary, zoomin, from_page, to_page, callback)
        callback(msg="OCR finished ({:.2f}s)".format(timer() - start))
        logging.debug("OCR: {}".format(timer() - start))

        start = timer()
        self._layouts_rec(zoomin)
        callback(0.65, "Layout analysis ({:.2f}s)".format(timer() - start))
        logging.debug("layouts: {}".format(timer() - start))

        start = timer()
        self._table_transformer_job(zoomin)
        callback(0.67, "Table analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._text_merge()
        tbls = self._extract_table_figure(True, zoomin, True, True)
        self._concat_downward()
        self._filter_forpages()
        callback(0.68, "Text merged ({:.2f}s)".format(timer() - start))

        # clean mess
        for b in self.boxes:
            b["text"] = re.sub(r"([\t 　]|\u3000){2,}", " ", b["text"].strip())

        return [(b["text"], b.get("layoutno", ""), self.get_position(b, zoomin)) for i, b in enumerate(self.boxes)], tbls


class Docx(DocxParser):
    def __init__(self):
        pass

    def get_picture(self, document, paragraph):
        img = paragraph._element.xpath(".//pic:pic")
        if not img:
            return None
        try:
            img = img[0]
            embed = img.xpath(".//a:blip/@r:embed")[0]
            related_part = document.part.related_parts[embed]
            image = related_part.image
            if image is not None:
                image = Image.open(BytesIO(image.blob))
                return image
            elif related_part.blob is not None:
                image = Image.open(BytesIO(related_part.blob))
                return image
            else:
                return None
        except Exception:
            return None

    def concat_img(self, img1, img2):
        if img1 and not img2:
            return img1
        if not img1 and img2:
            return img2
        if not img1 and not img2:
            return None
        width1, height1 = img1.size
        width2, height2 = img2.size

        new_width = max(width1, width2)
        new_height = height1 + height2
        new_image = Image.new("RGB", (new_width, new_height))

        new_image.paste(img1, (0, 0))
        new_image.paste(img2, (0, height1))

        return new_image

    def __call__(self, filename, binary=None, from_page=0, to_page=100000, callback=None):
        self.doc = Document(filename) if not binary else Document(BytesIO(binary))
        pn = 0
        last_answer, last_image = "", None
        question_stack, level_stack = [], []
        ti_list = []
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            question_level, p_text = 0, ""
            if from_page <= pn < to_page and p.text.strip():
                question_level, p_text = docx_question_level(p)
            if not question_level or question_level > 6:  # not a question
                last_answer = f"{last_answer}\n{p_text}"
                current_image = self.get_picture(self.doc, p)
                last_image = self.concat_img(last_image, current_image)
            else:  # is a question
                if last_answer or last_image:
                    sum_question = "\n".join(question_stack)
                    if sum_question:
                        ti_list.append((f"{sum_question}\n{last_answer}", last_image))
                    last_answer, last_image = "", None

                i = question_level
                while question_stack and i <= level_stack[-1]:
                    question_stack.pop()
                    level_stack.pop()
                question_stack.append(p_text)
                level_stack.append(question_level)
            for run in p.runs:
                if "lastRenderedPageBreak" in run._element.xml:
                    pn += 1
                    continue
                if "w:br" in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        if last_answer:
            sum_question = "\n".join(question_stack)
            if sum_question:
                ti_list.append((f"{sum_question}\n{last_answer}", last_image))

        tbls = []
        for tb in self.doc.tables:
            html = "<table>"
            for r in tb.rows:
                html += "<tr>"
                i = 0
                while i < len(r.cells):
                    span = 1
                    c = r.cells[i]
                    for j in range(i + 1, len(r.cells)):
                        if c.text == r.cells[j].text:
                            span += 1
                            i = j
                        else:
                            break
                    i += 1
                    html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                html += "</tr>"
            html += "</table>"
            tbls.append(((None, html), ""))
        return ti_list, tbls


def chunk(filename, binary=None, from_page=0, to_page=100000, lang="Chinese", callback=None, **kwargs):
    """
    Only pdf is supported.
    """
    parser_config = kwargs.get("parser_config", {"chunk_token_num": 512, "delimiter": "\n!?。；！？", "layout_recognize": "DeepDOC"})
    pdf_parser = None
    doc = {"docnm_kwd": filename}
    doc["title_tks"] = rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", doc["docnm_kwd"]))
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    # is it English
    eng = lang.lower() == "english"  # pdf_parser.is_english
    if re.search(r"\.pdf$", filename, re.IGNORECASE):
        layout_recognizer = parser_config.get("layout_recognize")
        if not layout_recognizer:
            # Fallback logic: Check if MinerU is configured in service_conf.yaml environment
            from common.constants import MINERU_DEFAULT_CONFIG, MINERU_ENV_KEYS
            mineru_conf = MINERU_DEFAULT_CONFIG.copy()
            import os
            
            # Check for MinerU token presence
            if os.environ.get("MINERU_TOKEN"):
                layout_recognizer = "MinerU"
                logging.info("Auto-selected MinerU based on MINERU_TOKEN env var")
            else:
                # Try reading service_conf.yaml directly if env var is missing
                try:
                    from common.config_utils import get_base_config
                    mineru_conf = get_base_config("mineru", {})
                    token = mineru_conf.get("token")
                    if token:
                        layout_recognizer = "MinerU"
                    else:
                        layout_recognizer = "DeepDOC"
                    logging.info(f"Auto-selection check: layout_recognizer={layout_recognizer} (token_in_env={bool(os.environ.get('MINERU_TOKEN'))}, token_in_conf={bool(token)})")
                except Exception as e:
                    logging.error(f"Failed to check service_conf for MinerU token: {e}")
                    layout_recognizer = "DeepDOC"
            # If user has explicitly enabled MinerU online mode, maybe they want it default?
            # But let's stick to safe defaults. DeepDOC is default.

        layout_recognizer, parser_model_name = normalize_layout_recognizer(layout_recognizer)

        if isinstance(layout_recognizer, bool):
            layout_recognizer = "DeepDOC" if layout_recognizer else "Plain Text"

        name = layout_recognizer.strip().strip('"').strip("'").lower()
        pdf_parser = PARSERS.get(name, by_plaintext)
        callback(0.1, "Start to parse.")

        kwargs.pop("parse_method", None)
        kwargs.pop("mineru_llm_name", None)
        sections, tbls, pdf_parser = pdf_parser(
            filename=filename,
            binary=binary,
            from_page=from_page,
            to_page=to_page,
            lang=lang,
            callback=callback,
            pdf_cls=Pdf,
            layout_recognizer=layout_recognizer,
            mineru_llm_name=parser_model_name,
            paddleocr_llm_name=parser_model_name,
            parse_method="manual",
            **kwargs,
        )

        if sections is None:
            sections = []
        if tbls is None:
            tbls = []

        # Check if sections are from MinerU with (text, level, poss) format
        is_mineru_format = False
        if sections and len(sections[0]) == 3:
            # MinerU format: (text, level, poss_list) where level is an int
            _, second, third = sections[0]
            if isinstance(second, int) and isinstance(third, list):
                is_mineru_format = True
                logging.info(f"[Manual] Detected MinerU format sections: {len(sections)} sections with levels")

        def _normalize_section(section):
            # pad section to length 3: (txt, sec_id, poss)
            if len(section) == 1:
                section = (section[0], "", [])
            elif len(section) == 2:
                section = (section[0], "", section[1])
            elif len(section) != 3:
                raise ValueError(f"Unexpected section length: {len(section)} (value={section!r})")

            txt, layoutno, poss = section
            
            # Handle MinerU format: poss is a list of position tuples
            if isinstance(poss, list) and poss and len(poss) > 0:
                # poss is [([page], left, right, top, bottom), ...]
                new_poss = []
                for p in poss:
                    if isinstance(p, tuple) and len(p) == 5:
                        pn, left, right, top, bottom = p
                        if isinstance(pn, list) and pn:
                            pn = pn[0]  # [page] -> page
                        new_poss.append((pn, left, right, top, bottom))
                poss = new_poss
            elif isinstance(poss, str):
                poss = pdf_parser.extract_positions(poss)
                if poss:
                    first = poss[0]  # tuple: ([pn], x1, x2, y1, y2)
                    pn = first[0]
                    if isinstance(pn, list) and pn:
                        pn = pn[0]  # [pn] -> pn
                        poss[0] = (pn, *first[1:])

            return (txt, layoutno, poss)

        sections = [_normalize_section(sec) for sec in sections]

        if not sections and not tbls:
            return []

        if name in ["tcadp", "docling", "mineru", "paddleocr"]:
            parser_config["chunk_token_num"] = 0
            logging.info(f"[Manual] Parser '{name}' detected, chunk_token_num set to 0 (trust parser's segmentation)")
        else:
            logging.info(f"[Manual] Parser '{name}' using chunk_token_num={parser_config.get('chunk_token_num', 512)}")

        callback(0.8, "Finish parsing.")

        # Use MinerU provided levels directly if available
        if is_mineru_format:
            logging.info(f"[Manual] Using MinerU provided text_level for section grouping")
            # sections are already (text, level, poss) from MinerU
            levels = [lvl for _, lvl, _ in sections]
            # Find the most common level (excluding 0 which means no specific level)
            from collections import Counter
            level_counts = Counter([l for l in levels if l > 0])
            if level_counts:
                most_level = level_counts.most_common(1)[0][0]
            else:
                most_level = 1
            logging.info(f"[Manual] MinerU levels: {levels}, most_level={most_level}")
        elif len(sections) > 0 and len(pdf_parser.outlines) / len(sections) > 0.03:
            max_lvl = max([lvl for _, lvl in pdf_parser.outlines])
            most_level = max(0, max_lvl - 1)
            levels = []
            for txt, _, _ in sections:
                for t, lvl in pdf_parser.outlines:
                    tks = set([t[i] + t[i + 1] for i in range(len(t) - 1)])
                    tks_ = set([txt[i] + txt[i + 1] for i in range(min(len(t), len(txt) - 1))])
                    if len(set(tks & tks_)) / max([len(tks), len(tks_), 1]) > 0.8:
                        levels.append(lvl)
                        break
                else:
                    levels.append(max_lvl + 1)

        else:
            bull = bullets_category([txt for txt, _, _ in sections])
            most_level, levels = title_frequency(bull, [(txt, lvl) for txt, lvl, _ in sections])

        assert len(sections) == len(levels)
        sec_ids = []
        sid = 0
        for i, lvl in enumerate(levels):
            if is_mineru_format:
                # For MinerU format, each section should be independent to preserve ordering
                # Don't merge sections with the same level
                sid += 1
            elif lvl <= most_level and i > 0 and lvl != levels[i - 1]:
                sid += 1
            sec_ids.append(sid)

        sections = [(txt, sec_ids[i], poss) for i, (txt, _, poss) in enumerate(sections)]
        logging.info(f"[Manual] Assigned sec_ids: {sec_ids[:10]}... (showing first 10)")
        logging.info(f"[Manual] Processing {len(tbls)} tables, from_page={from_page}")
        # Note: Tables are NOT added to sections here to avoid duplication.
        # They will be processed separately by tokenize_table and merged later.

        def tag(pn, left, right, top, bottom):
            if pn + left + right + top + bottom == 0:
                return ""
            # pn is 0-indexed page_idx, add 1 to make it 1-indexed for display
            return "@@{}\t{:.1f}\t{:.1f}\t{:.1f}\t{:.1f}##".format(pn + 1, left, right, top, bottom)

        chunks = []
        last_sid = -2
        tk_cnt = 0
        # For MinerU format, don't sort sections to preserve original order from JSON
        # This ensures tables and text are processed in the correct sequence
        if is_mineru_format:
            logging.info(f"[Manual] Preserving original section order for MinerU format (no sorting)")
            sections_to_process = sections
        else:
            sections_to_process = sorted(sections, key=lambda x: (x[-1][0][0], x[-1][0][3], x[-1][0][1]))
        for txt, sec_id, poss in sections_to_process:
            poss = "\t".join([tag(*pos) for pos in poss])
            # For MinerU format, don't merge sections to preserve each section as independent chunk
            if is_mineru_format:
                chunks.append(txt + poss)
                tk_cnt = num_tokens_from_string(txt)
                last_sid = sec_id
            elif tk_cnt < 32 or (tk_cnt < 1024 and (sec_id == last_sid or sec_id == -1)):
                if chunks:
                    chunks[-1] += "\n" + txt + poss
                    tk_cnt += num_tokens_from_string(txt)
                    continue
                chunks.append(txt + poss)
                tk_cnt = num_tokens_from_string(txt)
                if sec_id > -1:
                    last_sid = sec_id
            else:
                chunks.append(txt + poss)
                tk_cnt = num_tokens_from_string(txt)
                if sec_id > -1:
                    last_sid = sec_id
        logging.info(f"[Manual] Generated {len(chunks)} text chunks from {len(sections)} sections (chunk_token_num={parser_config.get('chunk_token_num', 512)})")
        logging.info(f"[Manual] Found {len(tbls)} tables to process")
        
        # Fix poss format in tbls: convert from [([page], left, right, top, bottom)] to [(page, left, right, top, bottom)]
        fixed_tbls = []
        for idx, ((img, rows), poss) in enumerate(tbls):
            logging.info(f"[Manual] Table {idx} before fix: img={img is not None}, rows_length={len(rows) if isinstance(rows, str) else 0}, poss={poss}")
            fixed_poss = []
            for p in poss:
                # p is ([page_idx], left, right, top, bottom) or (page_idx, left, right, top, bottom)
                if isinstance(p[0], list):
                    fixed_p = (p[0][0], p[1], p[2], p[3], p[4])
                    logging.info(f"[Manual] Table {idx}: converted poss from {p} to {fixed_p}")
                else:
                    fixed_p = p
                    logging.info(f"[Manual] Table {idx}: poss already in correct format: {fixed_p}")
                fixed_poss.append(fixed_p)
            fixed_tbls.append(((img, rows), fixed_poss))
            logging.info(f"[Manual] Table {idx} after fix: fixed_poss={fixed_poss}")
        tbls = fixed_tbls
        logging.info(f"[Manual] Fixed {len(tbls)} tables poss format")
        
        tbls = vision_figure_parser_pdf_wrapper(
            tbls=tbls,
            sections=sections,
            callback=callback,
            **kwargs,
        )
        
        # Tokenize tables and text chunks separately, then merge by position
        table_chunks = tokenize_table(tbls, doc, eng)
        text_chunks = tokenize_chunks(chunks, doc, eng, pdf_parser)
        logging.info(f"[Manual] Tokenized {len(table_chunks)} table chunks and {len(text_chunks)} text chunks")
        
        # Log ALL table chunks for debugging table splitting issues
        logging.info(f"[Manual] ===== TABLE CHUNKS DEBUG START =====")
        for idx, tc in enumerate(table_chunks):
            pos = tc.get('position_int', [])
            content = tc.get('content_with_weight', '')
            doc_type = tc.get('doc_type_kwd', 'unknown')
            # Extract caption if present
            caption_match = ""
            if "<caption>" in content:
                cap_match = re.search(r'<caption>(.*?)</caption>', content)
                if cap_match:
                    caption_match = cap_match.group(1)[:50]
            logging.info(f"[Manual] Table chunk {idx}: type={doc_type}, page={pos[0][0] if pos else 'N/A'}, top={pos[0][3] if pos else 'N/A'}, caption='{caption_match}', content_length={len(content)}")
        logging.info(f"[Manual] ===== TABLE CHUNKS DEBUG END =====")
        
        # Log first few text chunks for debugging
        for idx, tc in enumerate(text_chunks[:5]):
            pos = tc.get('position_int', [])
            content = tc.get('content_with_weight', '')[:80]
            logging.info(f"[Manual] Text chunk {idx}: position_int={pos}, content={content}...")
        
        # Merge table_chunks and text_chunks by position (page, top)
        # Each chunk has 'position_int' field with positions
        def get_chunk_position(item):
            """Get (page, top, original_index) position for sorting"""
            chunk, original_idx = item
            if 'position_int' in chunk and chunk['position_int']:
                # position_int is a list of [page, left, right, top, bottom]
                first_pos = chunk['position_int'][0]
                result = (first_pos[0], first_pos[3], original_idx)  # (page, top, original_index)
                logging.info(f"[Manual][get_chunk_position] chunk content={chunk.get('content_with_weight', '')[:30]}... -> position={result}")
                return result
            logging.info(f"[Manual][get_chunk_position] chunk has no position_int, returning (0, 0, {original_idx})")
            return (0, 0, original_idx)
        
        # Combine all chunks with original index to preserve order when positions are equal
        all_chunks_with_index = [(chunk, idx) for idx, chunk in enumerate(table_chunks + text_chunks)]
        logging.info(f"[Manual] Combining {len(all_chunks_with_index)} chunks before sorting")
        
        # Sort by position (page, top, original_index)
        all_chunks_with_index.sort(key=get_chunk_position)
        
        # Extract chunks without index
        all_chunks = [chunk for chunk, _ in all_chunks_with_index]
        
        # Log sorted chunks
        for idx, chunk in enumerate(all_chunks[:5]):
            pos = chunk.get('position_int', [])
            content = chunk.get('content_with_weight', '')[:50]
            doc_type = chunk.get('doc_type_kwd', 'unknown')
            logging.info(f"[Manual] Sorted chunk {idx}: type={doc_type}, position_int={pos}, content={content}...")
        
        res = all_chunks
        logging.info(f"[Manual] Final total chunks: {len(res)} (tables + text, sorted by position)")
        table_ctx = max(0, int(parser_config.get("table_context_size", 0) or 0))
        image_ctx = max(0, int(parser_config.get("image_context_size", 0) or 0))
        if table_ctx or image_ctx:
            attach_media_context(res, table_ctx, image_ctx)
        return res

    elif re.search(r"\.docx?$", filename, re.IGNORECASE):
        docx_parser = Docx()
        ti_list, tbls = docx_parser(filename, binary, from_page=0, to_page=10000, callback=callback)
        tbls = vision_figure_parser_docx_wrapper(sections=ti_list, tbls=tbls, callback=callback, **kwargs)
        res = tokenize_table(tbls, doc, eng)
        for text, image in ti_list:
            d = copy.deepcopy(doc)
            if image:
                d["image"] = image
                d["doc_type_kwd"] = "image"
            tokenize(d, text, eng)
            res.append(d)
        table_ctx = max(0, int(parser_config.get("table_context_size", 0) or 0))
        image_ctx = max(0, int(parser_config.get("image_context_size", 0) or 0))
        if table_ctx or image_ctx:
            attach_media_context(res, table_ctx, image_ctx)
        return res
    else:
        raise NotImplementedError("file type not supported yet(pdf and docx supported)")


if __name__ == "__main__":
    import sys

    def dummy(prog=None, msg=""):
        pass

    chunk(sys.argv[1], callback=dummy)
